"""
Document Processor
Phase 2.1: Backend Infrastructure & RAG Pipeline

Processes documents using lightweight libraries that work fully offline.
No HuggingFace model downloads required.

Supported formats:
- PDF: Using pypdf (pure Python, no external dependencies)
- DOCX: Using python-docx
- TXT, MD, HTML: Direct text processing
"""

import tempfile
from pathlib import Path
from typing import Any

import structlog

logger = structlog.get_logger()


class DocumentProcessor:
    """Process documents using lightweight offline-capable libraries."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".html"}

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
    ):
        """
        Initialize the document processor.

        Args:
            chunk_size: Target size for text chunks (in characters)
            chunk_overlap: Overlap between consecutive chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def _chunk_text(self, text: str) -> list[str]:
        """
        Split text into overlapping chunks.

        Args:
            text: Text to chunk

        Returns:
            List of text chunks
        """
        if not text:
            return []

        chunks = []
        start = 0
        text_length = len(text)

        while start < text_length:
            end = start + self.chunk_size

            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence endings
                for boundary in [". ", ".\n", "! ", "!\n", "? ", "?\n"]:
                    last_boundary = text.rfind(boundary, start, end)
                    if last_boundary > start:
                        end = last_boundary + len(boundary)
                        break

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self.chunk_overlap
            if start >= text_length:
                break

        return chunks

    def _extract_pdf_text(self, file_path: Path) -> tuple[str, int]:
        """
        Extract text from PDF using pypdf.

        Args:
            file_path: Path to PDF file

        Returns:
            Tuple of (extracted text, page count)

        Raises:
            ValueError: If PDF is encrypted, corrupted, or contains no extractable text
            ImportError: If pypdf is not installed
        """
        import traceback

        try:
            from pypdf import PdfReader
            from pypdf.errors import (
                PdfReadError,
                FileNotDecryptedError,
                EmptyFileError,
                PdfStreamError,
            )
        except ImportError as e:
            logger.error(
                "pypdf_not_installed",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ImportError(
                "pypdf is required for PDF processing. Install with: pip install pypdf"
            ) from e

        # Attempt to open and read the PDF file
        try:
            reader = PdfReader(str(file_path))
        except EmptyFileError as e:
            logger.error(
                "pdf_empty_file",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                "PDF file is empty or has zero bytes. Please ensure the file was uploaded correctly."
            ) from e
        except FileNotDecryptedError as e:
            logger.error(
                "pdf_encrypted_not_decrypted",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                "PDF is password-protected and could not be decrypted. "
                "Please provide an unencrypted version or the correct password."
            ) from e
        except PdfStreamError as e:
            logger.error(
                "pdf_stream_error",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            # PdfStreamError typically indicates truncated or malformed PDF
            error_msg = str(e).lower()
            if "truncated" in error_msg or "ended unexpectedly" in error_msg:
                raise ValueError(
                    "PDF file is incomplete or truncated. The file may not have been fully downloaded or saved. "
                    "Try re-downloading or re-saving the PDF."
                ) from e
            else:
                raise ValueError(
                    f"PDF file structure is invalid or corrupted: {e}. "
                    "Try opening and re-saving the PDF in a PDF reader application."
                ) from e
        except PdfReadError as e:
            logger.error(
                "pdf_read_error",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            # Provide specific guidance based on error message
            error_msg = str(e).lower()
            if "xref" in error_msg or "trailer" in error_msg:
                raise ValueError(
                    "PDF file appears to be corrupted (invalid cross-reference table). "
                    "Try opening and re-saving the PDF in a PDF reader application."
                ) from e
            elif "header" in error_msg:
                raise ValueError(
                    "PDF file has an invalid header. The file may be corrupted or not a valid PDF."
                ) from e
            elif "eof" in error_msg or "end" in error_msg:
                raise ValueError(
                    "PDF file is incomplete or truncated. The file may not have been fully downloaded or saved."
                ) from e
            else:
                raise ValueError(
                    f"Failed to read PDF file due to corruption or format issues: {e}"
                ) from e
        except PermissionError as e:
            logger.error(
                "pdf_permission_denied",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                "Permission denied when accessing the PDF file. "
                "Please ensure the file is not locked by another application."
            ) from e
        except FileNotFoundError as e:
            logger.error(
                "pdf_file_not_found",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"PDF file not found at path: {file_path}. The file may have been moved or deleted."
            ) from e
        except OSError as e:
            logger.error(
                "pdf_os_error",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"System error when accessing PDF file: {e}. "
                "Check file permissions and disk space."
            ) from e
        except MemoryError as e:
            logger.error(
                "pdf_memory_error",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                "PDF file is too large to process in available memory. "
                "Try splitting the PDF into smaller files."
            ) from e
        except Exception as e:
            # Catch-all for unexpected errors with full traceback
            error_type = type(e).__name__
            logger.error(
                "pdf_read_failed_unexpected",
                path=str(file_path),
                error_type=error_type,
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Unexpected error reading PDF file ({error_type}): {e}. "
                "The file may be in an unsupported format or corrupted."
            ) from e

        # Check if PDF is encrypted (additional check after successful read)
        try:
            if reader.is_encrypted:
                logger.error(
                    "pdf_encrypted",
                    path=str(file_path),
                    message="Encrypted PDFs require a password and are not supported",
                )
                raise ValueError(
                    "PDF is password-protected. Please provide an unencrypted version of the document."
                )
        except Exception as e:
            # Handle cases where is_encrypted check itself fails
            logger.error(
                "pdf_encryption_check_failed",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            # Continue processing - encryption check failure shouldn't block valid PDFs

        page_count = len(reader.pages)
        if page_count == 0:
            logger.warning("pdf_empty", path=str(file_path))
            raise ValueError("PDF has no pages")

        text_parts = []
        pages_with_text = 0
        extraction_errors = []

        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_parts.append(f"--- Page {page_num} ---\n{page_text}")
                    pages_with_text += 1
                else:
                    logger.debug(
                        "pdf_page_no_text",
                        path=str(file_path),
                        page=page_num,
                        message="Page may be scanned image or contain no text",
                    )
            except KeyError as e:
                error_msg = f"Page {page_num}: Missing required PDF object - {e}"
                extraction_errors.append(error_msg)
                logger.warning(
                    "pdf_page_missing_object",
                    path=str(file_path),
                    page=page_num,
                    error=str(e),
                    traceback=traceback.format_exc(),
                )
            except AttributeError as e:
                error_msg = f"Page {page_num}: Invalid page structure - {e}"
                extraction_errors.append(error_msg)
                logger.warning(
                    "pdf_page_invalid_structure",
                    path=str(file_path),
                    page=page_num,
                    error=str(e),
                    traceback=traceback.format_exc(),
                )
            except Exception as e:
                error_msg = f"Page {page_num}: {type(e).__name__} - {e}"
                extraction_errors.append(error_msg)
                logger.warning(
                    "pdf_page_extraction_failed",
                    path=str(file_path),
                    page=page_num,
                    error_type=type(e).__name__,
                    error=str(e),
                    traceback=traceback.format_exc(),
                )

        # Log extraction errors summary if any
        if extraction_errors:
            logger.warning(
                "pdf_extraction_partial_failure",
                path=str(file_path),
                total_pages=page_count,
                pages_extracted=pages_with_text,
                failed_pages=len(extraction_errors),
                errors=extraction_errors[:5],  # Log first 5 errors
            )

        full_text = "\n\n".join(text_parts)

        # Check if any text was extracted
        if not full_text.strip():
            logger.error(
                "pdf_no_extractable_text",
                path=str(file_path),
                page_count=page_count,
                message="PDF may contain only scanned images. OCR is required for such documents.",
            )
            raise ValueError(
                f"No text could be extracted from PDF ({page_count} pages). "
                "The document may contain only scanned images. OCR is required for such documents."
            )

        logger.info(
            "pdf_extracted",
            path=str(file_path),
            pages=page_count,
            pages_with_text=pages_with_text,
            chars=len(full_text),
            extraction_errors=len(extraction_errors),
        )
        return full_text, page_count

    def _extract_docx_text(self, file_path: Path) -> tuple[str, int | None]:
        """
        Extract text from DOCX using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Tuple of (extracted text, None for page count)

        Raises:
            ValueError: If DOCX is corrupted or cannot be read
            ImportError: If python-docx is not installed
        """
        import traceback

        try:
            from docx import Document as DocxDocument
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError as e:
            logger.error(
                "docx_not_installed",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ImportError(
                "python-docx is required for DOCX processing. Install with: pip install python-docx"
            ) from e

        try:
            doc = DocxDocument(str(file_path))
        except PackageNotFoundError as e:
            logger.error(
                "docx_not_found",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                "DOCX file not found or is not a valid Office Open XML document. "
                "The file may be corrupted or in an unsupported format."
            ) from e
        except KeyError as e:
            logger.error(
                "docx_missing_part",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                "DOCX file is missing required components. The file may be corrupted. "
                "Try opening and re-saving the document in Microsoft Word or LibreOffice."
            ) from e
        except ValueError as e:
            logger.error(
                "docx_invalid_format",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Invalid DOCX file format: {e}. "
                "Ensure the file is a valid .docx (not .doc) document."
            ) from e
        except PermissionError as e:
            logger.error(
                "docx_permission_denied",
                path=str(file_path),
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                "Permission denied when accessing the DOCX file. "
                "Please ensure the file is not locked by another application."
            ) from e
        except Exception as e:
            error_type = type(e).__name__
            logger.error(
                "docx_read_failed_unexpected",
                path=str(file_path),
                error_type=error_type,
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Unexpected error reading DOCX file ({error_type}): {e}. "
                "The file may be in an unsupported format or corrupted."
            ) from e

        try:
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                logger.warning(
                    "docx_no_text",
                    path=str(file_path),
                    message="DOCX contains no extractable text",
                )
                raise ValueError(
                    "DOCX file contains no extractable text. The document may be empty or contain only images."
                )

            logger.info(
                "docx_extracted", path=str(file_path), paragraphs=len(text_parts), chars=len(full_text)
            )
            return full_text, None
        except ValueError:
            # Re-raise ValueError from empty document check
            raise
        except Exception as e:
            error_type = type(e).__name__
            logger.error(
                "docx_extraction_failed",
                path=str(file_path),
                error_type=error_type,
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Error extracting text from DOCX ({error_type}): {e}"
            ) from e

    def _extract_text_file(self, file_path: Path) -> tuple[str, None]:
        """
        Read plain text files (txt, md, html).

        Args:
            file_path: Path to text file

        Returns:
            Tuple of (file content, None for page count)

        Raises:
            ValueError: If file cannot be decoded or read
        """
        import traceback

        # Try common encodings
        encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]
        decode_errors = []

        for encoding in encodings:
            try:
                with open(file_path, encoding=encoding) as f:
                    content = f.read()
                logger.info(
                    "text_file_read", path=str(file_path), encoding=encoding, chars=len(content)
                )
                return content, None
            except UnicodeDecodeError as e:
                decode_errors.append(f"{encoding}: {e}")
                continue
            except FileNotFoundError as e:
                logger.error(
                    "text_file_not_found",
                    path=str(file_path),
                    error=str(e),
                    traceback=traceback.format_exc(),
                )
                raise ValueError(
                    f"Text file not found at path: {file_path}. The file may have been moved or deleted."
                ) from e
            except PermissionError as e:
                logger.error(
                    "text_file_permission_denied",
                    path=str(file_path),
                    error=str(e),
                    traceback=traceback.format_exc(),
                )
                raise ValueError(
                    "Permission denied when accessing the text file. "
                    "Please ensure the file is not locked by another application."
                ) from e
            except OSError as e:
                logger.error(
                    "text_file_os_error",
                    path=str(file_path),
                    error=str(e),
                    traceback=traceback.format_exc(),
                )
                raise ValueError(
                    f"System error when accessing text file: {e}"
                ) from e

        # If we get here, none of the encodings worked
        logger.error(
            "text_file_decode_failed",
            path=str(file_path),
            tried_encodings=encodings,
            errors=decode_errors,
        )
        raise ValueError(
            f"Could not decode file {file_path.name} with any supported encoding "
            f"(tried: {', '.join(encodings)}). The file may be binary or use an unsupported encoding."
        )

    def _strip_html_tags(self, text: str) -> str:
        """
        Remove HTML tags from text (basic implementation).

        Args:
            text: HTML content

        Returns:
            Text with HTML tags removed
        """
        import re

        # Remove script and style elements
        text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)

        # Remove HTML tags
        text = re.sub(r"<[^>]+>", " ", text)

        # Clean up whitespace
        text = re.sub(r"\s+", " ", text)

        return text.strip()

    async def process_file(self, file_path: Path) -> dict[str, Any]:
        """
        Process a file and return chunks.

        Args:
            file_path: Path to the file

        Returns:
            Dictionary with chunks, metadata, and full text

        Raises:
            ValueError: If file type is not supported or processing fails
        """
        import traceback

        # Check if file exists before trying to get size
        if not file_path.exists():
            logger.error(
                "file_not_found",
                path=str(file_path),
            )
            raise ValueError(
                f"File not found at path: {file_path}. The file may have been moved or deleted."
            )

        logger.info("processing_file", path=str(file_path), size_bytes=file_path.stat().st_size)

        suffix = file_path.suffix.lower()
        if suffix not in self.SUPPORTED_EXTENSIONS:
            logger.error(
                "unsupported_file_type",
                path=str(file_path),
                extension=suffix,
                supported=list(self.SUPPORTED_EXTENSIONS),
            )
            raise ValueError(
                f"Unsupported file type: {suffix}. "
                f"Supported formats: {', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        # Extract text based on file type
        page_count = None

        try:
            if suffix == ".pdf":
                full_text, page_count = self._extract_pdf_text(file_path)
            elif suffix == ".docx":
                full_text, page_count = self._extract_docx_text(file_path)
            elif suffix == ".html":
                raw_text, _ = self._extract_text_file(file_path)
                full_text = self._strip_html_tags(raw_text)
            else:
                # .txt, .md - read as plain text
                full_text, _ = self._extract_text_file(file_path)
        except (ValueError, ImportError):
            # Re-raise known exceptions from extraction methods
            raise
        except Exception as e:
            # Catch any unexpected errors during extraction
            error_type = type(e).__name__
            logger.error(
                "file_extraction_failed",
                path=str(file_path),
                file_type=suffix,
                error_type=error_type,
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Failed to extract text from {suffix} file ({error_type}): {e}"
            ) from e

        # Chunk text
        try:
            chunks = self._chunk_text(full_text)
            if not chunks:
                logger.warning(
                    "no_chunks_generated",
                    path=str(file_path),
                    text_length=len(full_text),
                )
                raise ValueError(
                    f"No text chunks could be generated from {file_path.name}. "
                    "The document may be empty or contain only whitespace."
                )
        except Exception as e:
            error_type = type(e).__name__
            logger.error(
                "chunking_failed",
                path=str(file_path),
                error_type=error_type,
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Failed to chunk text from {file_path.name} ({error_type}): {e}"
            ) from e

        # Build metadata
        metadata = {
            "filename": file_path.name,
            "extension": suffix,
            "page_count": page_count,
        }

        logger.info(
            "file_processed",
            path=str(file_path),
            chunk_count=len(chunks),
            total_chars=len(full_text),
            pages=page_count,
        )

        return {
            "chunks": chunks,
            "metadata": metadata,
            "full_text": full_text,
        }

    async def process_bytes(
        self,
        content: bytes,
        filename: str,
    ) -> dict[str, Any]:
        """
        Process file content from bytes.

        Args:
            content: File content as bytes
            filename: Original filename (used for extension detection)

        Returns:
            Dictionary with chunks, metadata, and full text

        Raises:
            ValueError: If processing fails or file type is unsupported
        """
        import traceback

        if not content:
            logger.error("empty_bytes_content", filename=filename)
            raise ValueError(
                f"File {filename} has no content (0 bytes). Please ensure the file was uploaded correctly."
            )

        suffix = Path(filename).suffix
        logger.info(
            "processing_bytes",
            filename=filename,
            size_bytes=len(content),
            extension=suffix,
        )

        tmp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                tmp.write(content)
                tmp_path = Path(tmp.name)

            result = await self.process_file(tmp_path)
            return result
        except (ValueError, ImportError):
            # Re-raise known exceptions
            raise
        except OSError as e:
            logger.error(
                "temp_file_creation_failed",
                filename=filename,
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Failed to create temporary file for processing {filename}: {e}"
            ) from e
        except Exception as e:
            error_type = type(e).__name__
            logger.error(
                "process_bytes_failed",
                filename=filename,
                error_type=error_type,
                error=str(e),
                traceback=traceback.format_exc(),
            )
            raise ValueError(
                f"Failed to process uploaded file {filename} ({error_type}): {e}"
            ) from e
        finally:
            # Clean up temporary file
            if tmp_path and tmp_path.exists():
                try:
                    tmp_path.unlink()
                except Exception as e:
                    logger.warning(
                        "temp_file_cleanup_failed",
                        path=str(tmp_path),
                        error=str(e),
                    )

    async def process_text(self, text: str, source_name: str = "text") -> dict[str, Any]:
        """
        Process plain text directly (no file conversion needed).

        Args:
            text: Text content
            source_name: Name for the source

        Returns:
            Dictionary with chunks and metadata
        """
        chunks = self._chunk_text(text)

        return {
            "chunks": chunks,
            "metadata": {
                "filename": source_name,
                "extension": ".txt",
                "page_count": None,
            },
            "full_text": text,
        }
